import os, re
from docx import Document

TOURS_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(TOURS_DIR, "output_utf8.txt")

DOCX_FILES = [
    "Болгария_От_Дуная_до_Босфора_сокращено.docx",
    "Варшава_Берлин_Познань_сокращено.docx",
    "ВДП (Варшава, Дрезден, Прага) сокращено.docx",
    "Франция_и_ее_соседи_сокращено.docx",
    "Французский_поцелуй_сокращено.docx",
]

URL_RE = re.compile(r"https?://[^\s]+")

def extract_hyperlinks(doc):
    links = []
    for rel_id, rel in doc.part.rels.items():
        if "hyperlink" in str(rel.reltype).lower():
            links.append(rel.target_ref)
    return links

def check_docx(filepath, out):
    filename = os.path.basename(filepath)
    out.write("=" * 70 + "\n")
    out.write("FILE: " + filename + "\n")
    out.write("=" * 70 + "\n")
    try:
        doc = Document(filepath)
    except Exception as e:
        out.write("  ERROR opening document: " + str(e) + "\n")
        return

    text_urls = []
    context_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        found = URL_RE.findall(text)
        if found:
            text_urls.extend(found)
        lower = text.lower()
        if "подробная информация" in lower or "ссылка на тур" in lower:
            context_paragraphs.append((i, text))

    hyperlinks = extract_hyperlinks(doc)

    table_urls = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    found = URL_RE.findall(para.text)
                    if found:
                        table_urls.extend(found)

    out.write("\n")
    out.write("  [1] URL in paragraph text:      " + ("YES" if text_urls else "NO") + "\n")
    if text_urls:
        for u in text_urls:
            out.write("       -> " + u + "\n")
    out.write("  [2] URL in table cells:         " + ("YES" if table_urls else "NO") + "\n")
    if table_urls:
        for u in table_urls:
            out.write("       -> " + u + "\n")
    out.write("  [3] Hyperlink relationships:    " + ("YES" if hyperlinks else "NO") + "\n")
    if hyperlinks:
        for h in hyperlinks:
            out.write("       -> " + h + "\n")

    out.write("\n")
    out.write("  [4] Context around PODROBNAYA INFO / SSYLKA NA TUR:\n")
    if context_paragraphs:
        for idx, txt in context_paragraphs:
            out.write("       Para #" + str(idx) + ": " + txt[:300] + "\n")
    else:
        out.write("       (none found)\n")

    found_any = bool(text_urls or table_urls or hyperlinks)
    out.write("\n")
    out.write("  >>> SUMMARY: " + ("URL / HYPERLINK FOUND" if found_any else "NO URL FOUND") + " <<<\n")
    out.write("\n")

def main():
    with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
        out.write("Scanning " + str(len(DOCX_FILES)) + " .docx files in " + TOURS_DIR + "\n")
        out.write("=" * 70 + "\n")
        out.write("\n")
        for fname in DOCX_FILES:
            fpath = os.path.join(TOURS_DIR, fname)
            if not os.path.isfile(fpath):
                out.write("WARNING: File not found -> " + fpath + "\n")
                continue
            check_docx(fpath, out)
        out.write("Done.\n")
    print("Output written to: " + OUTPUT_FILE)

if __name__ == "__main__":
    main()
