from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(r"D:\projects\travel-agent-bot\tours\programma_tura_измененный.docx")
OUTPUT = Path(
    r"D:\projects\travel-agent-bot\tours\programma_tura_измененный_актуальный.docx"
)


def find_exact(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph, found {len(matches)}: {text!r}")
    return matches[0]


def find_nth(document: Document, text: str, occurrence: int):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) <= occurrence:
        raise RuntimeError(
            f"Expected occurrence {occurrence + 1}, found {len(matches)}: {text!r}"
        )
    return matches[occurrence]


def replace_once(paragraph, old: str, new: str) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if full_text.count(old) != 1:
        raise RuntimeError(
            f"Expected one replacement in paragraph, found {full_text.count(old)}: {old!r}"
        )

    start = full_text.index(old)
    end = start + len(old)
    offsets = []
    cursor = 0
    for run in paragraph.runs:
        offsets.append((cursor, cursor + len(run.text)))
        cursor += len(run.text)

    first = next(i for i, (_, run_end) in enumerate(offsets) if run_end > start)
    last = next(i for i, (run_start, _) in enumerate(offsets) if run_start < end <= offsets[i][1])
    first_start, _ = offsets[first]
    last_start, _ = offsets[last]
    prefix = paragraph.runs[first].text[: start - first_start]
    suffix = paragraph.runs[last].text[end - last_start :]

    if first == last:
        paragraph.runs[first].text = prefix + new + suffix
        return

    paragraph.runs[first].text = prefix + new
    for index in range(first + 1, last):
        paragraph.runs[index].text = ""
    paragraph.runs[last].text = suffix


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_date_before(reference, template, text: str) -> None:
    paragraph = reference.insert_paragraph_before()
    if paragraph._p.pPr is not None:
        paragraph._p.remove(paragraph._p.pPr)
    if template._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    run = paragraph.add_run(text)
    if template.runs and template.runs[0]._r.rPr is not None:
        run._r.insert(0, deepcopy(template.runs[0]._r.rPr))


def main() -> None:
    if OUTPUT.exists():
        raise FileExistsError(f"Refusing to overwrite existing output: {OUTPUT}")

    document = Document(SOURCE)

    # ВДП: remove the expired departure while retaining the "Когда:" label.
    first_vdp_date = find_exact(
        document, "Когда: 01.07.2026 - 05.07.2026 (12 мест)"
    )
    first_vdp_date.runs[0].text = "Когда: 05.08.2026 - 09.08.2026"
    remove_paragraph(find_exact(document, "05.08.2026 - 09.08.2026"))

    vdp_warsaw = find_exact(
        document,
        "2 день (первый день программы, четверг): Прибытие в Варшаву. "
        "обзорная экскурсия. Рыночную площадь с символом города – русалочкой, "
        "кафедральный собор св. Иоанна, королевский замок (без посещения "
        "интерьеров), Краковское предместье с дворцами польской аристократии "
        "и другие значимые достопримечательности Старого города. Переезд в "
        "транзитный отель у границы с Германией на ночлег.",
    )
    vdp_warsaw.add_run(
        " Внимание! В случае существенной задержки при прохождении границы "
        "возможен перенос экскурсии по Варшаве на воскресенье."
    )

    # Франция и её соседи: new 2027 departures and material itinerary notes.
    france_route = find_exact(
        document,
        "Куда: Минск - Бамберг - Ротенбург на Таубере - Страсбург - "
        "Баден-Баден - Люцерн - Лаутербруннен - Мюррен - Дижон- "
        "Кольмар-Риквир - Нюрнберг - Минск",
    )
    for city in ("Баден-Баден", "Люцерн", "Лаутербруннен", "Мюррен", "Дижон"):
        replace_once(france_route, city, city + "*")

    france_first_date = find_exact(document, "Когда: 12.12.2026 - 21.12.2026")
    france_cost = find_exact(
        document,
        "Сколько стоит: 630 € на человека при проживании в 2 или 3 местном номере",
    )
    for date in (
        "20.03.2027 - 28.03.2027",
        "08.05.2027 - 16.05.2027",
        "25.09.2027 - 03.10.2027",
        "11.12.2027 - 19.12.2027",
    ):
        insert_date_before(france_cost, france_first_date, date)

    france_murren = find_exact(
        document,
        "5 день, среда: Завтрак в отеле. Свободный день в Мюлузе либо "
        "выездная экскурсия по желанию за доплату в Люцерн- "
        "Лаутербруннен+Мюррен*. Возвращение в Мюлуз (или его пригород). "
        "Ночлег в отеле.",
    )
    replace_once(
        france_murren,
        "Возвращение в Мюлуз (или его пригород).",
        "В случае ревизионных работ или оползней канатная дорога в Мюррен "
        "может не работать; тогда проводится альтернативная экскурсионная "
        "программа с гидом по маршруту Лаутербруннен - Гриндельвальд - "
        "Интерлакен. Возвращение в Мюлуз (или его пригород).",
    )

    france_dijon = find_exact(
        document,
        "6 день, четверг: Завтрак в отеле. Свободный день в Мюлузе. Для "
        "желающих выезд в столицу Бургундии - Дижон. Возвращение в Мюлуз "
        "( Или его пригород). Ночлег в отеле.",
    )
    replace_once(
        france_dijon,
        "Для желающих выезд в столицу Бургундии - Дижон.",
        "Для желающих выезд в столицу Бургундии - Дижон* для "
        "самостоятельной прогулки-квеста.",
    )

    # Варшава — Берлин — Познань: factual clarifications, no date changes.
    berlin_cost = find_exact(
        document,
        "Сколько стоит: 220 € (на человека при проживании в 2-местном номере)",
    )
    replace_once(berlin_cost, "220 € (", "220 € без скрытых доплат! (")

    berlin_transfer = find_nth(
        document,
        "Внимание! в случае перегруженности границы и возникновении "
        "необходимости довоза в ожидающий автобус, стоимость тура "
        "увеличивается на 20 €",
        1,
    )
    replace_once(
        berlin_transfer,
        "довоза в ожидающий автобус",
        "довоза в ожидающий в очереди автобус",
    )

    berlin_headphones = find_nth(document, "наушники по программе", 1)
    berlin_headphones.runs[0].text = (
        "использование наушников во все экскурсионные дни"
    )

    berlin_return = find_exact(
        document,
        "5 день (день приезда, воскресенье): Прибытие домой в первой половине дня.",
    )
    replace_once(
        berlin_return,
        "Прибытие домой в первой половине дня.",
        "Прибытие домой рано утром.",
    )

    # Французский поцелуй: remove the expired date, update availability,
    # and append the 2027 schedule.
    french_first_date = find_exact(
        document, "Когда: 20.06.2026 - 28.06.2026 (мест нет)"
    )
    french_first_date.runs[0].text = (
        "Когда: 18.07.2026 - 26.07.2026 (2 места для туристов с визами)"
    )
    remove_paragraph(
        find_exact(
            document, "18.07.2026 - 26.07.2026 (2 места для туристов с визами)"
        )
    )

    september = find_exact(document, "12.09.2026 - 20.09.2026 (мест нет)")
    september.runs[0].text = (
        "12.09.2026 - 20.09.2026 (2 места для туристов с визами)"
    )
    october = find_exact(
        document, "24.10.2026 - 01.11.2026 (осталось 2 места)"
    )
    october.runs[0].text = "24.10.2026 - 01.11.2026 (мест нет)"

    french_template = find_exact(document, "19.12.2026 - 27.12.2026")
    french_cost = find_exact(
        document,
        "Сколько стоит: 630 € (на человека при проживании в 2 или 3 местном номере)",
    )
    for date in (
        "13.02.2027 - 21.02.2027",
        "06.03.2027 - 14.03.2027",
        "08.05.2027 - 16.05.2027",
        "12.06.2027 - 20.06.2027",
        "10.07.2027 - 18.07.2027",
        "07.08.2027 - 15.08.2027",
        "18.09.2027 - 26.09.2027",
        "16.10.2027 - 24.10.2027",
        "06.11.2027 - 14.11.2027",
        "04.12.2027 - 12.12.2027",
    ):
        insert_date_before(french_cost, french_template, date)

    document.save(OUTPUT)

    verified = Document(OUTPUT)
    if len(verified.paragraphs) != 219:
        raise RuntimeError(
            f"Unexpected paragraph count: {len(verified.paragraphs)} (expected 219)"
        )

    expected_once = (
        "Когда: 05.08.2026 - 09.08.2026",
        "20.03.2027 - 28.03.2027",
        "11.12.2027 - 19.12.2027",
        "Сколько стоит: 220 € без скрытых доплат! "
        "(на человека при проживании в 2-местном номере)",
        "Когда: 18.07.2026 - 26.07.2026 (2 места для туристов с визами)",
        "12.09.2026 - 20.09.2026 (2 места для туристов с визами)",
        "24.10.2026 - 01.11.2026 (мест нет)",
        "04.12.2027 - 12.12.2027",
    )
    texts = [paragraph.text for paragraph in verified.paragraphs]
    for expected in expected_once:
        if texts.count(expected) != 1:
            raise RuntimeError(f"Validation failed for {expected!r}")

    forbidden = (
        "Когда: 01.07.2026 - 05.07.2026 (12 мест)",
        "Когда: 20.06.2026 - 28.06.2026 (мест нет)",
        "12.09.2026 - 20.09.2026 (мест нет)",
        "24.10.2026 - 01.11.2026 (осталось 2 места)",
    )
    for stale in forbidden:
        if stale in texts:
            raise RuntimeError(f"Stale text remains: {stale!r}")

    print(f"Created: {OUTPUT}")
    print(f"Paragraphs: {len(verified.paragraphs)}")


if __name__ == "__main__":
    main()
